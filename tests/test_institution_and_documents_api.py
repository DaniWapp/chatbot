"""Pruebas HTTP de institución (nombre/logo dinámico) y documentos (subida,
listado, recategorización y borrado con etiqueta de dependencia).
embed_texts y vector_store se simulan para no cargar el modelo real de
embeddings ni tocar el índice FAISS real -- mismo patrón que
tests/test_ingest_service.py."""
from unittest.mock import patch

from fastapi.testclient import TestClient

from app.main import app
from app.services import admin_service

client = TestClient(app)

_counter = 0


def _fake_embed_texts(texts):
    return [[0.1, 0.2, 0.3] for _ in texts]


def _login_as(role="root", password="clave-segura-123", dependencia_id=None):
    global _counter
    _counter += 1
    username = f"test-inst-doc-{role}-{_counter}"
    admin_service.create_admin(username, password, f"Admin {role}", role, dependencia_id)
    res = client.post("/api/auth/login", json={"username": username, "password": password})
    return res.json()["token"]


def _auth(token):
    return {"Authorization": f"Bearer {token}"}


def test_get_institution_has_sensible_defaults_when_never_configured():
    res = client.get("/api/institution")

    assert res.status_code == 200
    body = res.json()
    assert body["name"]  # algún nombre por defecto, aunque nadie lo haya configurado
    assert body["logo_url"] is None


def test_update_institution_requires_root():
    token = _login_as(role="general")

    res = client.put(
        "/api/root/institution",
        data={"name": "Universidad Nueva", "extra_info": ""},
        headers=_auth(token),
    )

    assert res.status_code == 403


def test_update_institution_name_reflected_in_public_endpoint():
    token = _login_as()

    update_res = client.put(
        "/api/root/institution",
        data={"name": "Universidad de Prueba", "extra_info": "Contacto: test@uni.edu"},
        headers=_auth(token),
    )
    assert update_res.status_code == 200
    assert update_res.json()["name"] == "Universidad de Prueba"

    public_res = client.get("/api/institution")
    assert public_res.json()["name"] == "Universidad de Prueba"
    assert public_res.json()["extra_info"] == "Contacto: test@uni.edu"


def test_document_upload_list_recategorize_and_delete(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    dep = client.post(
        "/api/root/dependencias",
        json={"name": "Dep Documentos Test", "description": "desc"},
        headers=_auth(token),
    ).json()

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        upload_res = client.post(
            "/api/root/documents",
            files={
                "file": (
                    "prueba.txt",
                    b"Contenido de prueba con suficiente longitud para generar un chunk.",
                    "text/plain",
                )
            },
            data={"dependencia_id": str(dep["id"])},
            headers=_auth(token),
        )
        assert upload_res.status_code == 200
        assert upload_res.json()["documents_processed"] == 1

        list_res = client.get("/api/root/documents", headers=_auth(token))
        assert list_res.status_code == 200
        docs = list_res.json()
        assert any(d["filename"] == "prueba.txt" and d["dependencia_id"] == dep["id"] for d in docs)

        recategorize_res = client.put(
            "/api/root/documents/prueba.txt", json={"dependencia_id": None}, headers=_auth(token)
        )
        assert recategorize_res.status_code == 200

        list_res_2 = client.get("/api/root/documents", headers=_auth(token))
        doc_after = next(d for d in list_res_2.json() if d["filename"] == "prueba.txt")
        assert doc_after["dependencia_id"] is None

        delete_res = client.delete("/api/root/documents/prueba.txt", headers=_auth(token))
        assert delete_res.status_code == 200
        assert not (tmp_path / "prueba.txt").exists()


def test_document_upload_rejects_disallowed_extension(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    res = client.post(
        "/api/root/documents",
        files={"file": ("prueba.exe", b"contenido", "application/octet-stream")},
        headers=_auth(token),
    )

    assert res.status_code == 400


# --- /api/admin/documents: general (paridad con root) y dependencia (solo lo suyo) ---


def _upload_via_panel(token, filename, content=b"Contenido de prueba con suficiente longitud.", dependencia_id=None):
    data = {}
    if dependencia_id is not None:
        data["dependencia_id"] = str(dependencia_id)
    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        return client.post(
            "/api/admin/documents",
            files={"file": (filename, content, "text/plain")},
            data=data,
            headers=_auth(token),
        )


def test_dependencia_admin_upload_ignores_submitted_dependencia_and_forces_own(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep = client.post(
        "/api/root/dependencias", json={"name": "Dep Panel Docs", "description": "desc"}, headers=_auth(root_token)
    ).json()
    other_dep = client.post(
        "/api/root/dependencias", json={"name": "Otra Dep Panel Docs", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_token = _login_as(role="dependencia", dependencia_id=dep["id"])

    # Intenta subirlo etiquetado a OTRA dependencia -- debe ignorarse y forzar la suya.
    res = _upload_via_panel(dep_token, "dep-doc.txt", dependencia_id=other_dep["id"])
    assert res.status_code == 200

    from app.services import ingest_service

    assert ingest_service.get_document_dependencia("dep-doc.txt") == dep["id"]


def test_dependencia_admin_sees_only_own_documents(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep = client.post(
        "/api/root/dependencias", json={"name": "Dep Visibilidad", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_token = _login_as(role="dependencia", dependencia_id=dep["id"])
    general_token = _login_as(role="general")

    _upload_via_panel(dep_token, "propio.txt")
    _upload_via_panel(general_token, "de-otro.txt")  # general/compartido, no de esta dependencia

    dep_list = client.get("/api/admin/documents", headers=_auth(dep_token)).json()
    filenames = {d["filename"] for d in dep_list}
    assert "propio.txt" in filenames
    assert "de-otro.txt" not in filenames

    general_list = client.get("/api/admin/documents", headers=_auth(general_token)).json()
    general_filenames = {d["filename"] for d in general_list}
    assert {"propio.txt", "de-otro.txt"}.issubset(general_filenames)


def test_dependencia_admin_cannot_delete_document_of_another_dependencia(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep_a = client.post(
        "/api/root/dependencias", json={"name": "Dep A Docs", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_b = client.post(
        "/api/root/dependencias", json={"name": "Dep B Docs", "description": "desc"}, headers=_auth(root_token)
    ).json()
    token_a = _login_as(role="dependencia", dependencia_id=dep_a["id"])
    token_b = _login_as(role="dependencia", dependencia_id=dep_b["id"])

    _upload_via_panel(token_a, "solo-a.txt")

    res = client.delete("/api/admin/documents/solo-a.txt", headers=_auth(token_b))
    assert res.status_code == 403

    res_owner = client.delete("/api/admin/documents/solo-a.txt", headers=_auth(token_a))
    assert res_owner.status_code == 200


def test_dependencia_admin_cannot_recategorize():
    root_token = _login_as()
    dep = client.post(
        "/api/root/dependencias", json={"name": "Dep No Recategoriza", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_token = _login_as(role="dependencia", dependencia_id=dep["id"])

    res = client.put(
        "/api/admin/documents/algo.txt", json={"dependencia_id": None}, headers=_auth(dep_token)
    )

    assert res.status_code == 403


def test_general_admin_can_recategorize_and_delete_any_document(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep = client.post(
        "/api/root/dependencias", json={"name": "Dep General Parity", "description": "desc"}, headers=_auth(root_token)
    ).json()
    general_token = _login_as(role="general")

    _upload_via_panel(general_token, "general-parity.txt", dependencia_id=dep["id"])

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        recategorize_res = client.put(
            "/api/admin/documents/general-parity.txt", json={"dependencia_id": None}, headers=_auth(general_token)
        )
    assert recategorize_res.status_code == 200

    delete_res = client.delete("/api/admin/documents/general-parity.txt", headers=_auth(general_token))
    assert delete_res.status_code == 200


def test_admin_documents_requires_authentication():
    res = client.get("/api/admin/documents")
    assert res.status_code == 401


# --- Vista previa del texto extraído -------------------------------------


def test_root_can_preview_document_text(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()
    (tmp_path / "preview.txt").write_text("Contenido de ejemplo para la vista previa.", encoding="utf-8")

    res = client.get("/api/root/documents/preview.txt/preview", headers=_auth(token))

    assert res.status_code == 200
    body = res.json()
    assert body["filename"] == "preview.txt"
    assert body["text"] == "Contenido de ejemplo para la vista previa."
    assert body["truncated"] is False


def test_preview_truncates_long_documents(tmp_path, monkeypatch):
    from app.config import settings
    from app.api import routes

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    monkeypatch.setattr(routes, "_PREVIEW_MAX_CHARS", 20)
    token = _login_as()
    (tmp_path / "largo.txt").write_text("A" * 100, encoding="utf-8")

    res = client.get("/api/root/documents/largo.txt/preview", headers=_auth(token))

    assert res.status_code == 200
    body = res.json()
    assert len(body["text"]) == 20
    assert body["truncated"] is True


def test_preview_missing_document_returns_404(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    res = client.get("/api/root/documents/no-existe.txt/preview", headers=_auth(token))

    assert res.status_code == 404


def test_dependencia_admin_cannot_preview_document_of_another_dependencia(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep_a = client.post(
        "/api/root/dependencias", json={"name": "Dep Preview A", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_b = client.post(
        "/api/root/dependencias", json={"name": "Dep Preview B", "description": "desc"}, headers=_auth(root_token)
    ).json()
    token_a = _login_as(role="dependencia", dependencia_id=dep_a["id"])
    token_b = _login_as(role="dependencia", dependencia_id=dep_b["id"])

    _upload_via_panel(token_a, "solo-a-preview.txt")

    res_denied = client.get("/api/admin/documents/solo-a-preview.txt/preview", headers=_auth(token_b))
    assert res_denied.status_code == 403

    res_ok = client.get("/api/admin/documents/solo-a-preview.txt/preview", headers=_auth(token_a))
    assert res_ok.status_code == 200


def test_general_admin_can_preview_any_document(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    root_token = _login_as()
    dep = client.post(
        "/api/root/dependencias", json={"name": "Dep Preview General", "description": "desc"}, headers=_auth(root_token)
    ).json()
    dep_token = _login_as(role="dependencia", dependencia_id=dep["id"])
    general_token = _login_as(role="general")

    _upload_via_panel(dep_token, "de-dependencia-preview.txt")

    res = client.get("/api/admin/documents/de-dependencia-preview.txt/preview", headers=_auth(general_token))
    assert res.status_code == 200


# --- Conversión automática de PDF/DOCX a TXT al subir, sin conservar el original ---


def test_upload_docx_converts_to_txt_and_discards_original(tmp_path, monkeypatch):
    """Extremo a extremo con un DOCX real (sin mockear la extracción) --
    confirma que solo queda el .txt, con el texto real del documento, y que
    el .docx original no se conserva en el servidor."""
    from docx import Document as DocxDocument
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    docx_bytes_path = tmp_path / "_source.docx"
    doc = DocxDocument()
    doc.add_paragraph("Contenido real de prueba para la conversión automática a texto plano.")
    doc.save(str(docx_bytes_path))
    content = docx_bytes_path.read_bytes()
    docx_bytes_path.unlink()  # no hace parte del DOCUMENTS_DIR, solo se usó para generar los bytes

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        res = client.post(
            "/api/root/documents",
            files={"file": ("Convertible.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers=_auth(token),
        )

    assert res.status_code == 200
    assert not (tmp_path / "Convertible.docx").exists()
    txt_path = tmp_path / "Convertible.txt"
    assert txt_path.exists()
    assert "Contenido real de prueba" in txt_path.read_text(encoding="utf-8")

    docs = client.get("/api/root/documents", headers=_auth(token)).json()
    filenames = {d["filename"] for d in docs}
    assert "Convertible.txt" in filenames
    assert "Convertible.docx" not in filenames


def test_upload_pdf_converts_to_txt_and_discards_original(tmp_path, monkeypatch):
    """La extracción de PDF ya está cubierta a fondo en test_document_loader.py
    -- aquí se mockea load_document para probar solo la lógica nueva de
    conversión/descarte del original en la ruta de subida."""
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    from app.rag.document_loader import LoadedDocument, PageText

    fake_document = LoadedDocument(
        filename="Reporte.pdf", pages=[PageText(page_number=1, text="Texto extraído del PDF de prueba.")]
    )

    with (
        patch("app.api.routes.load_document", return_value=fake_document),
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        res = client.post(
            "/api/root/documents",
            files={"file": ("Reporte.pdf", b"%PDF-1.4 contenido falso, la extraccion esta mockeada", "application/pdf")},
            headers=_auth(token),
        )

    assert res.status_code == 200
    assert not (tmp_path / "Reporte.pdf").exists()
    txt_path = tmp_path / "Reporte.txt"
    assert txt_path.exists()
    assert txt_path.read_text(encoding="utf-8") == "Texto extraído del PDF de prueba."


def test_upload_pdf_extraction_failure_cleans_up_original(tmp_path, monkeypatch):
    from app.config import settings
    from app.rag.document_loader import DocumentLoadError

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    with patch("app.api.routes.load_document", side_effect=DocumentLoadError("PDF corrupto de prueba")):
        res = client.post(
            "/api/root/documents",
            files={"file": ("Corrupto.pdf", b"no es un pdf real", "application/pdf")},
            headers=_auth(token),
        )

    assert res.status_code == 400
    assert not (tmp_path / "Corrupto.pdf").exists()
    assert not (tmp_path / "Corrupto.txt").exists()


def test_upload_xlsx_is_not_converted_to_txt(tmp_path, monkeypatch):
    """XLSX se excluye a propósito de la conversión (ver
    docs/notas-mejora-documentos.md): perdería el chunking por fila."""
    from openpyxl import Workbook
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    xlsx_source = tmp_path / "_source.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["Materia", "Día"])
    ws.append(["Cálculo", "Lunes"])
    wb.save(str(xlsx_source))
    content = xlsx_source.read_bytes()
    xlsx_source.unlink()

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        res = client.post(
            "/api/root/documents",
            files={"file": ("Horario.xlsx", content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            headers=_auth(token),
        )

    assert res.status_code == 200
    assert (tmp_path / "Horario.xlsx").exists()
    assert not (tmp_path / "Horario.txt").exists()


def test_upload_txt_is_unaffected_by_conversion(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        res = client.post(
            "/api/root/documents",
            files={"file": ("normal.txt", b"contenido de prueba suficientemente largo.", "text/plain")},
            headers=_auth(token),
        )

    assert res.status_code == 200
    assert (tmp_path / "normal.txt").exists()


def test_upload_returns_final_filename_in_response(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    with (
        patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts),
        patch("app.services.ingest_service.vector_store.add_chunks"),
        patch("app.services.ingest_service.vector_store.reset_collection"),
    ):
        res = client.post(
            "/api/root/documents",
            files={"file": ("simple.txt", b"contenido de prueba suficientemente largo.", "text/plain")},
            headers=_auth(token),
        )

    assert res.status_code == 200
    assert res.json()["final_filename"] == "simple.txt"


def test_colliding_pdf_and_docx_get_consecutive_names_instead_of_overwriting(tmp_path, monkeypatch):
    """Reporte.pdf y Reporte.docx, subidos por separado, convergerían al
    mismo "Reporte.txt" -- el segundo no debe pisar al primero."""
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    from app.rag.document_loader import LoadedDocument, PageText

    with patch("app.services.ingest_service.embed_texts", side_effect=_fake_embed_texts), patch(
        "app.services.ingest_service.vector_store.add_chunks"
    ), patch("app.services.ingest_service.vector_store.reset_collection"):
        with patch(
            "app.api.routes.load_document",
            return_value=LoadedDocument(filename="Reporte.pdf", pages=[PageText(page_number=1, text="Contenido del PDF.")]),
        ):
            res1 = client.post(
                "/api/root/documents",
                files={"file": ("Reporte.pdf", b"contenido falso, extraccion mockeada", "application/pdf")},
                headers=_auth(token),
            )
        assert res1.status_code == 200
        assert res1.json()["final_filename"] == "Reporte.txt"

        with patch(
            "app.api.routes.load_document",
            return_value=LoadedDocument(filename="Reporte.docx", pages=[PageText(page_number=1, text="Contenido del DOCX.")]),
        ):
            res2 = client.post(
                "/api/root/documents",
                files={
                    "file": (
                        "Reporte.docx",
                        b"contenido falso, extraccion mockeada",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=_auth(token),
            )
        assert res2.status_code == 200
        assert res2.json()["final_filename"] == "Reporte (2).txt"

    # Ambos archivos siguen existiendo, con su contenido propio -- ninguno se perdió.
    assert (tmp_path / "Reporte.txt").read_text(encoding="utf-8") == "Contenido del PDF."
    assert (tmp_path / "Reporte (2).txt").read_text(encoding="utf-8") == "Contenido del DOCX."


def test_documents_list_is_sorted_newest_first(tmp_path, monkeypatch):
    import time
    from app.config import settings

    monkeypatch.setattr(settings, "DOCUMENTS_DIR", tmp_path)
    token = _login_as()

    (tmp_path / "primero.txt").write_text("contenido", encoding="utf-8")
    time.sleep(0.05)
    (tmp_path / "segundo.txt").write_text("contenido", encoding="utf-8")

    docs = client.get("/api/root/documents", headers=_auth(token)).json()
    filenames = [d["filename"] for d in docs]

    assert filenames.index("segundo.txt") < filenames.index("primero.txt")
