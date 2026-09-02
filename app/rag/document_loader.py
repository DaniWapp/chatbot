"""Carga y extracción de texto de documentos (PDF, DOCX, TXT, XLSX).

Cada documento se descompone en una lista de "páginas" de texto para poder
conservar la referencia de página en los metadatos de cada chunk. Los TXT y
DOCX no tienen paginación real, así que se tratan como una única "página".
Los XLSX usan una "página" por hoja de cálculo.
"""
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import settings


class DocumentLoadError(Exception):
    """Se lanza cuando un documento no puede leerse (corrupto, formato inválido, etc.)."""


@dataclass
class PageText:
    page_number: int  # 1-indexado; None-like sentinel = 1 para formatos sin páginas
    text: str


@dataclass
class LoadedDocument:
    filename: str
    pages: List[PageText]
    # True para hojas de cálculo: el chunker las trata distinto (empaqueta
    # filas completas en vez de cortar por cantidad de caracteres), para no
    # partir una fila de datos a la mitad ni diluirla mezclada con texto
    # de otro tema.
    is_tabular: bool = False


def validate_file(path: Path) -> None:
    if path.suffix.lower() not in settings.ALLOWED_EXTENSIONS:
        raise DocumentLoadError(
            f"Extensión no soportada: {path.suffix}. Permitidas: {settings.ALLOWED_EXTENSIONS}"
        )
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > settings.MAX_FILE_SIZE_MB:
        raise DocumentLoadError(
            f"Archivo demasiado grande ({size_mb:.1f}MB). Límite: {settings.MAX_FILE_SIZE_MB}MB"
        )
    if path.stat().st_size == 0:
        raise DocumentLoadError("El archivo está vacío.")


def _load_pdf(path: Path) -> List[PageText]:
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"No se pudo abrir el PDF (¿está corrupto?): {exc}") from exc

    pages: List[PageText] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(PageText(page_number=i, text=text))

    if not pages:
        raise DocumentLoadError("El PDF no contiene texto extraíble (¿está escaneado como imagen?).")
    return pages


def _load_docx(path: Path) -> List[PageText]:
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"No se pudo abrir el DOCX (¿está corrupto?): {exc}") from exc

    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text.strip():
        raise DocumentLoadError("El DOCX no contiene texto.")
    return [PageText(page_number=1, text=text)]


def _load_txt(path: Path) -> List[PageText]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="latin-1")
        except Exception as exc:
            raise DocumentLoadError(f"No se pudo leer el TXT: {exc}") from exc
    except Exception as exc:
        raise DocumentLoadError(f"No se pudo leer el TXT: {exc}") from exc

    if not text.strip():
        raise DocumentLoadError("El TXT está vacío.")
    return [PageText(page_number=1, text=text)]


def _load_xlsx(path: Path) -> List[PageText]:
    try:
        workbook = load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        raise DocumentLoadError(
            f"No se pudo abrir el Excel (¿está corrupto o es .xls antiguo? "
            f"guárdalo como .xlsx): {exc}"
        ) from exc

    pages: List[PageText] = []
    for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        header = [str(cell).strip() if cell is not None else "" for cell in rows[0]]
        row_lines: List[str] = []
        for row in rows[1:]:
            if all(cell is None for cell in row):
                continue
            parts = [
                f"{col_name}: {value}"
                for col_name, value in zip(header, row)
                if value is not None and str(value).strip() != ""
            ]
            if parts:
                row_lines.append(" | ".join(parts))

        if row_lines:
            pages.append(PageText(page_number=sheet_index, text="\n".join(row_lines)))

    if not pages:
        raise DocumentLoadError("El Excel no contiene datos legibles (¿hojas vacías?).")
    return pages


def load_document(path: Path) -> LoadedDocument:
    """Carga un documento y devuelve su texto dividido por página.

    Lanza DocumentLoadError con un mensaje claro si el archivo no es válido
    o está corrupto, para que el proceso de ingestión pueda omitirlo sin
    detener el resto del lote.
    """
    validate_file(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = _load_pdf(path)
    elif suffix == ".docx":
        pages = _load_docx(path)
    elif suffix == ".txt":
        pages = _load_txt(path)
    elif suffix == ".xlsx":
        pages = _load_xlsx(path)
        return LoadedDocument(filename=path.name, pages=pages, is_tabular=True)
    else:
        raise DocumentLoadError(f"Extensión no soportada: {suffix}")

    return LoadedDocument(filename=path.name, pages=pages)
